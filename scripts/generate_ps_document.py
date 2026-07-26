from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "PS_Jewellers_Complete_Website_Admin_Legal_Documentation.docx"
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(.7); section.bottom_margin = Inches(.7); section.left_margin = Inches(.8); section.right_margin = Inches(.8)
styles = doc.styles
styles["Normal"].font.name = "Aptos"; styles["Normal"].font.size = Pt(9.5)
for name, size, color in [("Title",28,"8B641E"),("Heading 1",18,"8B641E"),("Heading 2",13,"24364B")]:
    s=styles[name]; s.font.name="Aptos"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
if "Small Note" not in styles:
    s=styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH); s.font.name="Aptos"; s.font.size=Pt(8); s.font.italic=True; s.font.color.rgb=RGBColor(90,90,90)
header=section.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; header.add_run("PS JEWELLERS  |  DIGITAL SHOWROOM OPERATIONS HANDBOOK").font.size=Pt(8)
footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; footer.add_run("Confidential working documentation  •  PS Jewellers  •  ")
field=OxmlElement("w:fldSimple"); field.set(qn("w:instr"),"PAGE"); footer._p.append(field)
def shade(cell, color="8B641E"):
    shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),color); cell._tc.get_or_add_tcPr().append(shd)
def table(headers, rows):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,v in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=str(v); shade(c)
        for r in c.paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v)
    tr=t.rows[0]._tr.get_or_add_trPr(); th=OxmlElement("w:tblHeader"); th.set(qn("w:val"),"true"); tr.append(th)
    doc.add_paragraph(); return t
def p(text="", style=None): doc.add_paragraph(text, style=style)
def h(text, level=1): doc.add_heading(text, level=level)
def bullets(xs):
    for x in xs: p(x,"List Bullet")
def numbered(xs):
    for x in xs: p(x,"List Number")

# Cover and contents
q=doc.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=q.add_run("PS JEWELLERS"); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(139,100,30)
q=doc.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=q.add_run("Digital Showroom, Admin, Security and Legal Documentation"); r.bold=True; r.font.size=Pt(26); r.font.color.rgb=RGBColor(36,54,75)
q=doc.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=q.add_run("Comprehensive editable project handbook"); r.italic=True; r.font.size=Pt(13)
table(["Document control","Value"],[("System","PS Jewellers digital jewellery showroom and enquiry platform"),("Location","Jhunjhunu, Rajasthan, India — verify final public address"),("Website","https://ps-jewellers-mudita-diamond-collect.vercel.app/"),("Repository","https://github.com/manglamtechnicalagency-crypto/PS_Jewellers_Mudita_Diamond_Collection_Jhunujhu"),("Prepared","26 July 2026"),("Status","Implementation reference, operational handbook and legal-content draft"),("Approval","Owner and qualified Indian legal counsel must approve publication content")])
h("Confidentiality and legal notice",2)
p("This document describes the implemented application and provides operational/legal drafting material. It is not legal advice, a valuation, pricing approval, privacy impact assessment or guarantee of third-party configuration. Verify all legal text, prices, tax claims, hallmark/certification claims, product facts, media rights and contact details before publication.","Small Note")
doc.add_page_break()
h("Contents")
for x in ["1. Executive summary and product definition","2. Business model and customer journey","3. Public website map","4. Catalogue, product and media model","5. Admin control panel handbook","6. Architecture and data flow","7. API and integration inventory","8. Security, privacy and operations","9. Legal publication drafts","10. QA, deployment and load evidence","11. Backup and incident response","12. Launch checklist and sign-off","Appendix A. Environment variables","Appendix B. Source map"]: p(x)
p("Use Microsoft Word's Navigation Pane with Heading 1–2 styles. Replace bracketed placeholders before publication.","Small Note")

# Product and customer model
h("1. Executive summary and product definition")
p("PS Jewellers is a premium jewellery showroom website for presenting gold, diamond and silver jewellery, collecting customer enquiries and handing qualified interest to the showroom team. The public website is enquiry-led rather than a checkout/payment marketplace.")
h("1.1 Current capabilities",2)
bullets(["Searchable published catalogue with product codes, names, categories, purity, weight, descriptions, availability and pricing state.","Cloudflare R2 images/videos with Supabase metadata and product-media relationships.","Protected admin dashboard for products, media, enquiries, settings, pricing, audit and catalogue operations.","Server-rendered SEO metadata, canonical URLs, JSON-LD, sitemap and robots controls.","Enquiry creation before WhatsApp/follow-up handoff, with validation, rate limiting and idempotency."])
h("1.2 Non-goals and business boundary",2)
bullets(["Not a payment gateway, online fulfilment system, stock reservation engine or legally binding sales contract.","A shortlist, form, WhatsApp message, phone call or email is an enquiry until the showroom confirms availability, final price, taxes, terms and delivery/collection.","The catalogue is not proof of live stock unless the owner confirms the inventory source.","Supabase is the publication source of truth; Cloudflare D1 is an optional mirror/export."])
h("1.3 Current implementation evidence",2)
table(["Area","Current state"],[("Stack","Next.js App Router, React 19, TypeScript strict, Tailwind CSS"),("Data/auth","Supabase PostgreSQL, RLS, Supabase Auth and server TOTP/AAL2"),("Media","Cloudflare R2 with Supabase media metadata and links"),("Catalogue","17 published seed products; 78 product-media links; 10 price-on-request products"),("Roles","super_admin, admin, editor, viewer"),("Quality","Lint, type-check, 50 unit/security tests, build and Playwright smoke tests")])

h("2. Business model and customer journey")
numbered(["Discover: visitor arrives from home, category, campaign, search or a shared product URL.","Evaluate: visitor reviews product imagery/video, purity, weight, code, description, availability and price state.","Shortlist: visitor saves products; unpriced products remain enquiry-only.","Enquire: server validates the customer form, creates an enquiry, applies rate-limit/idempotency controls and creates the follow-up handoff.","Follow up: showroom staff assign, qualify, schedule a visit, negotiate and record outcome in CRM.","Confirm offline: showroom confirms stock, final quote, charges, tax, certification, delivery/collection and payment terms."])
h("2.1 Jewellery pricing policy",2)
p("Jewellery prices can depend on metal rate, purity, weight, wastage, making charges, stone/carats, certification, GST, discount and date. Separate approved, indicative and enquiry-required states. Record rate source, effective date, reason, approver and affected-product count for changes.")
table(["State","Customer behavior","Admin requirement"],[("Published price","Show approved amount and note","Record source/date/reason"),("Price on request","Hide numeric price; show enquiry CTA","Capture product intent"),("Draft/review","Do not expose in catalogue/sitemap","Complete facts/media/review"),("Archived/deleted","404 or omit","Retain audit/history where required")])

h("3. Public website map")
table(["Route group","Routes / purpose"],[("Core","/; /shop; /product/:slug"),("Collections","/gold-jewellery; /diamond-jewellery; /silver-jewellery; /bridal-collection; /rings; /necklaces; /earrings; /bangles; /bracelets; /chains; /pendants; /mangalsutra; /maang-tikka; /nose-pin; /anklets"),("Campaigns","/new-arrivals; /best-sellers; /offers"),("Customer tools","/wishlist; /cart; /checkout; /account; /order-tracking"),("Content/trust","/about; /contact; /faq; /blog; /store-locator; /book-appointment; /privacy-policy; /terms; /return-policy"),("SEO/system","/sitemap.xml; /robots.txt")])
h("3.1 Content requirements",2)
bullets(["Each product needs real name, stable SKU/slug, primary image, gallery, optional video, facts, availability and approved price/enquiry behavior.","Informative images need meaningful alt text; decorative images need intentional empty alt values.","Policy pages need owner-approved legal copy, effective date, business identity, contact and jurisdiction.","Contact details, showroom hours, address, map, WhatsApp and email must be verified.","Avoid unsupported claims such as certified, hallmarked, in stock, limited, best or discount."])

h("4. Catalogue, product and media model")
table(["Field group","Required information"],[("Identity","UUID, SKU, name, slug, category, collection, tags"),("Jewellery facts","Metal, purity, weight, stone/diamond details, dimensions, hallmark/certification"),("Commercial","Price, price-on-request flag, currency, rate source, making/wastage inputs"),("Publishing","Draft/review/scheduled/published, visibility, publish time, archive/delete"),("Availability","Stock/reserved/sold state, low-stock and customer message"),("Content","Short/full description, care, delivery/returns, SEO title/description"),("Relationships","Primary/gallery/video media, reviews, enquiries, audit history")])
h("4.1 R2 media lifecycle",2)
bullets(["Authenticate and authorize admin; validate MIME, size, path and rate limit before presigning.","Upload to an R2 quarantine key; register metadata only after validation/processing.","Store storage key, public URL, content type, size, hash/status, alt text, caption, usage and approval.","Link media to product with primary/gallery/video role and persisted order.","Replacement updates the new storage key atomically; remove the old object only after approval and reference checks.","Deletion blocks or unlinks product references before deleting an R2 object."])
h("4.2 Media checklist",2)
table(["Item","Acceptance rule"],[("Primary image","Exactly one approved primary; valid object and alt text"),("Gallery","Ordered image list; drag/drop order persisted; no orphan links"),("Video","Approved video with poster/thumbnail; CSP allows R2 media"),("Quality","Dimensions, size, duplicate hash, magic-byte/malware scan, variants"),("Governance","Quarantine/review/published/rejected/deleted; actor/timestamp audit")])

h("5. Admin control panel handbook")
table(["Role","Recommended scope"],[("super_admin","Roles, settings, pricing approvals, destructive actions, audit and recovery"),("admin","Products, media, enquiries, settings and operations"),("editor","Product/media/content editing and review submission"),("viewer","Read-only dashboards, catalogue, enquiries and audit")])
p("Every admin request must enforce session validity, verified MFA/AAL2 in production, profile role and database RLS. UI hiding is not authorization.")
table(["Route","Purpose"],[("/admin","Dashboard: products, active media, product-linked media, enquiries, audits and provider health"),("/admin/products","Search/filter/create/list, bulk operations, import and duplication"),("/admin/products/:id","Full product edit, pricing, media checklist, preview, publish and history"),("/admin/media","R2 gallery, product links, replacement, ordering, guards and pagination"),("/admin/enquiries","Search/filter, assignment, statuses, timeline, follow-up and export"),("/admin/audit","Bounded audit records with actor/date/entity filters"),("/admin/settings","Showroom identity, validated WhatsApp and public settings"),("/admin/catalogue","Catalogue sync/status controls"),("/admin/login","Password sign-in plus mandatory TOTP")])
h("5.1 Product publish workflow",2)
numbered(["Create draft with stable SKU/slug and taxonomy.","Add verified facts, price/enquiry state, description, SEO and availability.","Complete primary/gallery/video media checklist.","Preview the real storefront product route.","Submit for review and approve with role permissions.","Publish transactionally; verify product page, sitemap and enquiry CTA.","Rollback the prior version with reason in audit history."])
h("5.2 Enquiry CRM statuses",2)
table(["Status","Meaning"],[("new","Unassigned enquiry"),("qualified","Customer intent/product fit confirmed"),("follow_up","Next contact and owner recorded"),("visit_booked","Appointment scheduled"),("negotiation","Price/product discussion active"),("won","Sale confirmed outside the system"),("lost","Closed with mandatory reason"),("spam","Abuse or irrelevant message")])

h("6. Architecture and data flow")
p("The project is a modular monolith inside Next.js. MVC and feature boundaries create seams for later microservice extraction without duplicating authentication, migrations and transactions now.")
table(["Layer","Implementation"],[("View","Next.js Server Components, React client components, Tailwind, storefront and admin"),("Controller","app/api/**/route.ts: auth, origin checks, parsing, service calls, stable errors"),("Service","src/server/features/**: product and future media/enquiry/pricing services"),("Model","Supabase tables/views/RPCs, Zod schemas, repositories and R2 metadata"),("Storage","Cloudflare R2 for binary images/videos"),("Mirror","Optional Cloudflare D1 catalogue export; never overrides Supabase")])
h("6.1 Data flows",2)
bullets(["Catalogue: request → published Supabase view → product/media joins → R2 URLs → server-rendered page.","Admin product: browser → protected route → Supabase session/MFA/role → Zod → service → repository/RPC → audit.","Media: admin → presign/rate limit → R2 quarantine → processing/approval → metadata/link → public visibility.","Enquiry: browser → origin/rate-limit/validation/idempotency → enquiry snapshot → CRM timeline → WhatsApp handoff."])
table(["System","Source-of-truth responsibility"],[("Supabase","Published catalogue, facts, links, enquiries, admin profiles, settings, audit and pricing history"),("R2","Binary image/video objects and variants"),("D1","Optional synchronized export/mirror"),("Browser","Temporary shortlist/cart/recent state; never auth or durable truth")])

h("7. API and integration inventory")
table(["Endpoint","Purpose / control"],[("GET /api/catalogue","Published catalogue; fails closed when DB unavailable; media validation"),("GET /api/public/settings","Public-safe showroom settings"),("POST /api/public/enquiries","Same-origin, rate-limit, Zod, UUID validation and idempotency"),("/api/admin/products","CRUD with MFA, role, RLS, validation and audit"),("POST /api/admin/products/import","CSV preview, taxonomy validation and atomic import"),("POST /api/admin/products/bulk","Pricing RPC, history, reason and safe recalculation"),("/api/admin/media","Metadata, replacement, linking, ordering and reference-safe deletion"),("POST /api/admin/media/presign","R2 presign with size/MIME/rate/quarantine controls"),("/api/admin/enquiries","CRM search, filters, pagination, assignment and timeline"),("/api/admin/audit","Bounded, cursor-paginated audit records"),("POST /api/admin/d1-sync","Protected Supabase-to-D1 mirror sync")])
h("7.1 External services",2)
bullets(["Supabase Auth/Postgres/RLS: identity, sessions, MFA, persistent data and policies.","Cloudflare R2: object storage and public/private object delivery.","Vercel: Next.js deployment and serverless runtime.","WhatsApp: external communication channel; destination must be validated server configuration.","Email, analytics, maps and other vendors: configure only after privacy/consent/processor review."])

h("8. Security, privacy and operations")
h("8.1 Implemented controls",2)
bullets(["Server-side admin auth and roles; production requires TOTP/AAL2.","Supabase RLS; service credentials stay server-only.","Same-origin validation, structured errors and boundary validation.","R2 server-only credentials, short-lived URLs and upload rate limiting.","CSP, frame protection, referrer policy, content-type controls, noindex admin and secret scanning.","Audit records for admin mutations and workflow actions."])
h("8.2 Residual risks",2)
table(["Risk","Treatment"],[("R2 video CSP","Add configured R2 origin to media-src and verify playback; current deployed review found R2 video requests blocked"),("Media safety","Add magic-byte inspection, malware scanning, quarantine, poster and variants"),("Distributed limits","Configure Upstash/edge limiter; process-local limits do not coordinate across instances"),("Legal content","Replace placeholders, dates, identity, retention and jurisdiction after approval"),("Backup","Schedule Supabase exports, R2 versioning/retention and restore drills"),("Admin E2E","Use isolated staging admin storage state; do not test destructive flows in production")])
h("8.3 Privacy inventory",2)
table(["Data","Purpose","Retention/access"],[("Name, phone, email","Respond to enquiry","CRM staff; [OWNER TO APPROVE]"),("Product snapshot","Preserve enquiry context","Enquiry/audit; [OWNER TO APPROVE]"),("Consent/time","Evidence of customer choice","Restricted audit; [OWNER TO APPROVE]"),("Admin/MFA metadata","Secure administration","Security/admin; [OWNER TO APPROVE]"),("IP/request metadata","Abuse prevention","Restricted logs; [OWNER TO APPROVE]")])

h("9. Legal publication drafts")
p("These editable drafts adapt the supplied Desert Haveli reference handbook to a jewellery showroom. They are not legal advice. Owner and qualified Indian legal counsel must complete every bracketed field before publication.")
h("9.1 Terms and conditions",2)
numbered(["Effective date: [DATE]. Operator/legal entity: [LEGAL NAME], [REGISTERED ADDRESS], [CONTACT].","The site presents jewellery information, imagery, videos and showroom enquiry channels. Content, prices, availability and specifications may change.","A shortlist, form, WhatsApp message, email or call is an enquiry only—not an order, payment, reservation, stock hold or binding sale—until confirmed by PS Jewellers in writing.","Displayed prices may be indicative or depend on metal rate, purity, weight, making charges, wastage, stones, GST, certification and delivery. The showroom confirms the final quotation.","Customers must provide accurate information and must not misuse the website, upload harmful content or attempt unauthorized access.","External WhatsApp, maps, email, hosting and storage providers have separate terms/privacy notices.","Branding, photography, video, copy and software may not be reused without permission. Governing law/jurisdiction: [COUNSEL TO COMPLETE]."])
h("9.2 Privacy notice",2)
numbered(["Controller: [LEGAL ENTITY / OWNER]. Effective date: [DATE].","We may process contact details, product selection, enquiry message, preferences, consent, security logs and limited technical data to answer enquiries, operate and secure the website, manage the catalogue and comply with law.","Providers may include Vercel, Supabase, Cloudflare R2, WhatsApp, email, analytics and other approved processors. Confirm regions, subprocessors and transfer safeguards.","Do not use enquiry data for unrelated marketing without a valid lawful basis and opt-out process.","Rights/contact and retention periods: [PRIVACY CONTACT], [RETENTION PERIODS TO COMPLETE]."])
h("9.3 Price, enquiry, return and media notice",2)
p("Price on request means the showroom will confirm applicable pricing and availability after reviewing an enquiry. Return, exchange, cancellation, custom-order, deposit, delivery and refund terms belong in the final quotation or invoice and must comply with applicable law. Publish only licensed/owned media and authentic moderated reviews. Hallmark, diamond, gemstone, certification and ethical-sourcing claims require documentary support.")

h("10. QA, deployment and load-test evidence")
table(["Check","Evidence"],[("Lint/type-check/build","Passes"),("Unit/security tests","50 passed"),("Browser suite","8 public/protection tests passed; 10 staging-admin tests skipped without storage state"),("Route crawl","46 sitemap pages and 46 internal links; zero failures"),("Responsive/accessibility smoke","Mobile/desktop overflow passed; buttons/images had usable markup"),("Production load smoke","164 public GET requests; 100% HTTP 200; sequential p95 1,358 ms; burst p95 2,036 ms")])
h("10.1 Release verification",2)
numbered(["Verify deployment commit and Supabase migrations match.","Check RLS, admin MFA/roles and R2 bucket policies.","Verify R2 URL/CDN, lifecycle/versioning, quarantine and scanner.","Open category, product, price-on-request and policy pages.","Fix and verify R2 video CSP before claiming video support.","Submit one staging enquiry; verify CRM, idempotency and WhatsApp handoff.","Run browser suite with isolated staging admin storage state.","Review Vercel/Supabase/R2 logs, latency and errors after deploy."])
h("11. Backup and incident response")
table(["Asset","Control"],[("Supabase","Scheduled export, migration retention and isolated restore test"),("R2","Versioning/retention, metadata inventory and restore test"),("Secrets","Secret manager only; rotate on exposure"),("Code/docs","Protected GitHub main, CI gates and tagged releases"),("Enquiries/audit","Restricted export, access control and approved retention")])
numbered(["Detect: capture time, URL, request ID, deployment and affected entity without secrets.","Contain: revoke sessions/tokens, disable affected upload/admin action, quarantine/unpublish content.","Investigate: preserve logs, audit, R2 keys, database records and deployment details.","Recover: rotate credentials, restore safe data and rerun migrations/tests.","Notify owner/providers/affected people where required; document corrective action."])

h("12. Launch checklist and sign-off")
for x in ["Legal identity, address, phone, WhatsApp, email and business identifiers verified.","Products, SKUs, purity, weights, stones, hallmark/certification and media rights verified.","Metal rates, making charges, wastage, GST, discounts and price-on-request rules approved.","Terms, privacy, return/exchange and enquiry copy approved by counsel.","R2 secrets rotated and server-only; bucket lifecycle/versioning/public policy verified.","R2 video CSP fixed and desktop/mobile video playback verified.","Supabase migrations, RLS, MFA, roles and recovery tested in staging.","Quarantine, magic-byte/malware scan, poster/variants and orphan cleanup enabled.","CRM assignment, statuses, timeline, idempotency, rate limits and WhatsApp tested.","Backup restore drill and incident contacts completed.","Route, link, responsive, accessibility, SEO, performance and load checks completed."]: p("☐ "+x)
table(["Role","Name","Signature/date"],[("Showroom owner","[TO COMPLETE]","____________________________"),("Technical owner","[TO COMPLETE]","____________________________"),("Legal/privacy reviewer","[TO COMPLETE]","____________________________"),("Operations lead","[TO COMPLETE]","____________________________")])

doc.add_page_break(); h("Appendix A. Environment variables")
p("Use the deployment secret manager or uncommitted .env.local. Never put secrets in this Word file or browser-exposed variables.")
table(["Variable","Visibility","Purpose"],[("NEXT_PUBLIC_SUPABASE_URL","Public","Supabase URL"),("NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY","Public","Browser-safe key"),("SUPABASE_SECRET_KEY","Server-only","Privileged operations"),("R2_ACCOUNT_ID / R2_ACCESS_KEY_ID / R2_SECRET_ACCESS_KEY","Server-only","R2 integration"),("R2_BUCKET_NAME","Server-only","R2 destination"),("R2_UPLOAD_ADMIN_TOKEN","Server-only","Presign authorization"),("NEXT_PUBLIC_R2_PUBLIC_URL","Public","R2/CDN read URL"),("NEXT_PUBLIC_SITE_URL","Public","Canonical/sitemap/OG URL"),("UPSTASH_REDIS_REST_URL / TOKEN","Server-only","Distributed rate limiting"),("CLOUDFLARE_ACCOUNT_ID / D1_DATABASE_ID / API_TOKEN","Server-only","Optional D1 sync"),("WHATSAPP_SHOWROOM_NUMBER","Server-only","Validated enquiry handoff")])
h("Appendix B. Source map")
table(["Area","Primary files"],[("Public runtime","app/[[...slug]]/page.tsx; src/storefront-pages/*; src/components/*"),("Catalogue","src/lib/catalogue-server.ts; src/lib/catalogue-data.ts; app/api/catalogue/route.ts"),("Admin","app/admin/*; app/api/admin/*"),("Enquiry","app/api/public/enquiries/route.ts; src/lib/storefront-enquiry.ts"),("R2","src/lib/r2-server.ts; app/api/admin/media/*; app/api/r2-presign/route.ts"),("Supabase","src/lib/supabase/*; supabase/migrations/*"),("D1","src/lib/cloudflare-d1.ts; app/api/admin/d1-sync/route.ts"),("Security","proxy.ts; src/lib/security-headers.ts; src/lib/admin-auth.ts; docs/security/*"),("Architecture","docs/architecture/*"),("QA","docs/qa/*; tests/*"),("Reference","Desert_Haveli_Complete_Website_Admin_Legal_Documentation.pdf — adapted to jewellery context")])
p("End of document — replace bracketed fields, obtain approvals and update this handbook whenever runtime, policy, vendor, schema or deployment changes.","Small Note")
doc.core_properties.title="PS Jewellers Complete Website, Admin and Legal Documentation"
doc.core_properties.subject="Digital jewellery showroom operational handbook"
doc.core_properties.author="Manglam Technical Agency"
doc.core_properties.keywords="PS Jewellers, jewellery showroom, admin, Supabase, Cloudflare R2, legal, operations"
OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT); print(OUT)
